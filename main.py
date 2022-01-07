# @author Cesar Ramirez
# @program Agra Quote Calculator

from error import *
import math
from order_class import Order
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_UNDERLINE
from datetime import date
from num2words import num2words

# Collect price per square foot
sqft_price = input("Enter price per square foot: $")
empty_response = empty_string(sqft_price)
sqft_price = empty_literal(empty_response, "Please enter price per square foot: $", sqft_price)
sqft_price = float(sqft_price)

# DIAMETER configuration
diameter_ft = input("Enter diameter of tank (ft): ")
empty_response = empty_string(diameter_ft)
diameter_ft = empty_literal(empty_response, "Please enter diameter of tank (ft): ", diameter_ft)
diameter_ft = float(diameter_ft)

diameter_inch = input("Enter any remaining inches: ")
empty_response = empty_string(diameter_inch)
diameter_inch = empty_literal(empty_response, "Please enter any remaining inches: ", diameter_inch)
diameter_inch = float(diameter_inch)
diameter_tank = converter(diameter_ft, diameter_inch)

# DEPTH configuration
depth_ft = input("Enter depth of tank (ft): ")
empty_response = empty_string(depth_ft)
depth_ft = empty_literal(empty_response, "Please enter depth of tank (ft): ", depth_ft)
depth_ft = float(depth_ft)

depth_inch = input("Enter any remaining inches: ")
empty_response = empty_string(depth_inch)
depth_inch = empty_literal(empty_response, "Please enter any remaining inches: ", depth_inch)
depth_inch = float(depth_inch)

depth_tank = converter(depth_ft, depth_inch)

print("")

# Create object
quote_order = Order(sqft_price, diameter_tank, depth_tank)

# Print square footage
Order.print_square_footage(quote_order)

# Prints out cost of single liner
print("\nQuote cost of liner: ${:,.2f}".format(Order.get_single_liner_cost(quote_order)))

# Prompt for customization loop
print("\nCustomize order below:\nType \'help\' for options\nType \'back\' for menu\n------------------------\n")

# Setup for customization loop
order_list = []
satisfied = False
discounted = False
customized_installation_cost = False
total_quote_cost = Order.get_single_liner_cost(quote_order)
tax_perc = 0

# Customization loop
while not satisfied:

    # Reset back button
    back_button = False

    # Get user command
    command = input("> ").lower()

    # Accidentally hit space
    if command == "" or len(command) == 1:
        print("\nPlease enter help if you need command list\n")
        continue

    # To finish the order
    if command == 'finish':
        satisfied = True
        print("\nOrder Completed")
        continue

    # GEO
    elif command[0] == 'g' and command[1] == 'e':
        geo_satisfied = False

        total_quote_cost += Order.get_single_layer_geo_cost(quote_order, "both")

        # Keep track of order
        order_list.append("geo")

        # Prints out information
        print("\nCost of geo added: ${:,.2f}".format(Order.get_single_layer_geo_cost(quote_order, "both")))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # J-BOLTS
    elif command[0] == 'j':
        # Accounts in total quote cost
        total_quote_cost += Order.get_jbolt_cost(quote_order)

        # Keep track of order
        order_list.append("jbolt")

        print("\nCost of j-bolts added: ${:,.2f}".format(Order.get_jbolt_cost(quote_order)))
        print("\nTotal quote cost: ${:,.2f}\n".format(total_quote_cost))

    # INSTALLATION
    elif command[0] == 'i' and command[1] == 'n':
        installation_cost = Order.calc_installation_cost(quote_order)
        # Returns zero if diameter > 110
        if installation_cost == 0:
            installation_cost = input("\nPlease enter installation cost: $")
            customized_installation_cost = True
            empty_response = empty_string(installation_cost)
            # Back button
            if not empty_response and (installation_cost[0] == 'b' or installation_cost[0] == 'B'):
                print("")
                continue
            else:
                installation_cost = empty_literal(empty_response, "Please enter installation cost: $",
                                                  installation_cost)
                if installation_cost[0] == 'b' or installation_cost[0] == 'B':
                    print("")
                    continue
                installation_cost = float(installation_cost)

        if not customized_installation_cost:
            print()

        # Figures cost of traveling
        miles_traveled = input("Enter number of miles being charged: ")
        empty_response = empty_string(miles_traveled)
        if not empty_response and (miles_traveled[0] == 'b' or miles_traveled[0] == 'B'):
            print("")
            continue
        else:
            miles_traveled = empty_literal(empty_response, "Please enter number of miles being charged: "
                                           , miles_traveled)
            if miles_traveled == empty_response:
                miles_traveled = empty_literal(empty_response, "Please enter number of miles being charged: "
                                               , miles_traveled)
            if miles_traveled[0] == 'b' or miles_traveled[0] == 'B':
                print("")
                continue
        miles_traveled = int(miles_traveled)

        travel_cost = miles_traveled * 450
        print('Standard travel costs ${:,}.'.format(travel_cost))
        modified_travel_cost = input("Would you like to change it (yes/no)? ").lower()
        empty_response = empty_string(modified_travel_cost)
        if empty_response:
            modified_travel_cost = empty_literal(empty_response, "Please enter yes or no: ", modified_travel_cost)
        if modified_travel_cost[0] == 'y':
            travel_cost = input("Enter travel cost: $")
            empty_response = empty_string(travel_cost)
            if empty_response:
                travel_cost = empty_literal(empty_response, "Please enter travel cost: $", travel_cost)
            travel_cost = float(travel_cost)

        # Calculates total cost
        total_installation_cost = installation_cost + travel_cost

        # Accounts costs
        total_quote_cost += total_installation_cost

        # Keep track of order
        order_list.append("installation")

        # Prints out info
        print("\nCost of installation added: ${:,.2f}".format(installation_cost))
        print("Cost of mileage & mobilization added: ${:,.2f}".format(travel_cost))
        print("\nTotal cost of installation package added: ${:,.2f}\n".format(total_installation_cost))

    # ADD liners
    elif command[0] == 'a' and command[1] == 'd':
        # Collect number of liners
        additional_liners = input("\nEnter number of liners you wish to add: ")
        if additional_liners[0] == 'b' or additional_liners[0] == 'B':
            print("")
            continue
        else:
            additional_liners = int(additional_liners)

        # Sets total liners for quote
        Order.set_total_liners(quote_order, additional_liners)

        # Calculate total cost
        total_quote_cost += Order.get_total_liners_cost(quote_order) - Order.get_single_liner_cost(quote_order)

        # Print out final info
        print("\nTotal cost of all liners: ${:,.2f}\n".format(Order.get_total_liners_cost(quote_order)))

    # DISCOUNT liner price
    elif command[0] == 'd' and command[1] == 'i':

        # Collect desired amount of discount
        discount_amount_percentage = input("\nEnter percentage you wish to discount liner price: ")
        if discount_amount_percentage[0] == 'b' or discount_amount_percentage[0] == 'B':
            print("")
            continue
        else:
            discount_amount_percentage = float(discount_amount_percentage)
        # Turn percentage into a number
        discount_amount_number = discount_amount_percentage / 100
        # Subtract previous total liner cost
        total_quote_cost -= Order.get_total_liners_cost(quote_order)
        # Set & add new, discounted liner cost
        Order.set_liner_discount(quote_order, discount_amount_percentage)
        total_quote_cost += Order.get_total_liners_cost(quote_order)

        # Prints out new info
        print("\nCost of new liner with " + str(discount_amount_percentage) +
              "% discount: ${:,.2f}".format(Order.get_single_liner_cost(quote_order)))
        if quote_order.total_liners > 1:
            print("Total cost of all liners: ${:,.2f}\n".format(Order.get_total_liners_cost(quote_order)))
        else:
            print()

    # Lifting hem
    elif command[0] == 'l' and command[1] == 'i':

        # Calculate total cost
        Order.set_lifting_hem_cost(quote_order)
        total_quote_cost += Order.get_lifting_hem_cost(quote_order)

        # Keep track of order
        order_list.append("lifting hem")

        # Print out final info
        print("\nCost of lifting hem added: ${:,.2f}\n".format(Order.get_lifting_hem_cost(quote_order)))

    # Batten strips
    elif command[0] == "b" and command[1] == 'a':

        # Keep track of order
        order_list.append("batten strip")

        # Calculate total quote cost
        total_quote_cost += Order.get_batten_strip_cost(quote_order)

        # Print out final info
        print("\nCost of batten strips added: ${:,.2f}".format(Order.get_batten_strip_cost(quote_order)))
        print("Total quote cost: ${:,.2f}\n".format(total_quote_cost))

    # Taxes
    elif command[0] == 't' and command[1] == "a":

        # Find percentage user wishes to tax
        tax_perc = input("\nEnter tax percentage: ")
        empty_response = empty_string(tax_perc)
        if not empty_response and (tax_perc[0] == 'b' or tax_perc[0] == 'B'):
            print("")
            continue
        else:
            tax_perc = empty_literal(empty_response, "Please enter tax percentage: ", tax_perc)
        tax_perc = float(tax_perc)

        print("\nPlease finish order for final tax calculation\n")


    # Help commands
    elif command == 'help':
        print("\nCustomizations available:"
              "\n--------------------------\nGeo\nJ-bolts\nInstallation\nLifting hem\nBatten strips\nTax"
              "\nAdd liner(s)\nDiscount liner\n\nTo complete order enter \'finish\'\n\n")

    # For user mistake
    else:
        print("\nPlease enter help if you need command list\n")

# DOCUMENT SECTION

# Needed for documentation
total_liners = Order.get_total_liners(quote_order)

# Today's date
today = date.today()
date_pretty = today.strftime("%B %d, %Y")

# Create document object
quote = docx.Document()

# HEADER for quote

# Create header paragraph for quote
header = quote.add_paragraph(date_pretty)
header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Fill out important info
header.add_run('\nNAME OF CONTACT')
header.add_run('\nCOMPANY NAME')
header.add_run('\nPHONE NUMBER')
header.add_run('\nZip Code: FILL OUT HERE')
header.add_run('\nCITY, STATE')

# Double spacing for header
header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# GENERAL BODY for quote

# Shows number of liners ordered
if total_liners > 1:
    body = quote.add_paragraph(num2words(total_liners).capitalize() + f'({total_liners}) liners fabricated from ENTER'
                                                                      f' MATERIAL NAME HERE')
else:
    body = quote.add_paragraph(num2words(total_liners).capitalize() + f'({total_liners}) liner fabricated from ENTER'
                                                                      f' MATERIAL NAME HERE')

# Prints out dimensions
body.add_run(f'\n\n{round(diameter_ft)}\'-{round(diameter_inch)}\" diameter X {round(depth_ft)}\'-{round(depth_inch)}'
             f'\" deep. ')

# Prints out customizations to general body
if len(order_list) > 0:

    body.add_run(f'Includes ')

    # Loops through order list to add customizations to general body
    for order in order_list:

        # Sneaks in 'and'
        if order_list[-1] == order and len(order_list) > 1:
            body.add_run(f'and ')

        # Geo
        if order[0] == 'g' and order[1] == 'e':
            body.add_run('16oz geotextile padding for the floor and 8oz geotextile padding for the sidewalls')
            if order_list.index(order) != len(order_list) - 1 and len(order_list) != 2:
                body.add_run(", ")

        # Batten strip
        elif order[0] == 'b' and order[1] == 'a':
            body.add_run(f'batten strips')
            if order_list.index(order) != len(order_list) - 1 and len(order_list) != 2:
                body.add_run(", ")

        # J-bolts
        elif order[0] == 'j' and order[1] == 'b':
            body.add_run(f'{num2words(Order.get_jbolt_number(quote_order))} '
                         f'({Order.get_jbolt_number(quote_order)}) j-bolts')
            if order_list.index(order) != len(order_list) - 1 and len(order_list) != 2:
                body.add_run(", ")

        # Lifting hem
        elif order[0] == 'l' and order[1] == 'i':
            body.add_run('lifting hem')
            if order_list.index(order) != len(order_list) - 1 and len(order_list) != 2:
                body.add_run(", ")

        # Installation
        elif order[0] == 'i' and order[1] == 'n':
            body.add_run(f'installation')
            if order_list.index(order) != len(order_list) - 1 and len(order_list) != 2:
                body.add_run(", ")

        # Sneaks in space
        if order_list.index(order) == 0 and len(order_list) == 2:
            body.add_run(" ")

        # Sneaks in '.'
        if order_list.index(order) == len(order_list) - 1:
            body.add_run(".")

# LINING SYSTEM section of quote


# Liner subsection of LINING SYSTEM

calculations_feet = quote.add_paragraph("\nBottom square footage:                         "
                                        "                                                      "
                                        "{:,}'".format(Order.get_actual_square_footage(quote_order, "floor")))
calculations_feet.add_run("\nSidewall square footage:                  "
                          "                                                           ")
square_footage_underline = calculations_feet.add_run("{:,}'".
                                                     format(Order.get_actual_square_footage(quote_order, "wall")))
actual_square_footage = Order.get_actual_square_footage(quote_order, "both")
five_percent = actual_square_footage * 0.05
square_footage = math.ceil(five_percent + actual_square_footage)
# If lifting hem is ordered prints out square footage
if Order.get_lifting_hem_cost(quote_order) != 0:
    lifting_hem_sqft = Order.get_liner_circumference(quote_order)
    actual_square_footage += lifting_hem_sqft
    # I have included lifting hem in 5 percent
    five_percent = actual_square_footage * 0.05
    square_footage = round(five_percent + actual_square_footage)

    calculations_feet.add_run("\nLifting hem:                                              "
                              "                                                         ")
    calculations_feet.add_run("{:,}'".format(round(Order.get_liner_circumference(quote_order)))).underline = True
else:
    square_footage_underline.underline = True

calculations_feet.add_run("\nSquare footage:                                                "
                          "                                              "
                          "{:,}'".format(round(actual_square_footage)))

calculations_feet.add_run("\n5%:                                                         "
                          "                                                              ")
calculations_feet.add_run("{:,}'".format(round(five_percent))).underline = True

# Prints out the total square footage
total_sqft_underline = calculations_feet.add_run("\nTotal square footage:                                 "
                                                 "                                                  "
                                                 "{:,}'".format(square_footage))
total_sqft_underline.underline = WD_UNDERLINE.SINGLE

# Prints out cost of material
calculations_feet.add_run("\nCost of material:                                       "
                          "                                                       "
                          "${:,.2f}".format(sqft_price))

# Double spaces the paragraph
calculations_feet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Extensions subsection of LINING SYSTEM

# Calculate liner cost
liner_cost = square_footage * sqft_price

# Prints out liner cost (not modified)
extensions_cost = quote.add_paragraph("\nLiner cost:                                            "
                                      "                                                        "
                                      "${:,.2f}".format(liner_cost))

# Prints out discount if there is one
if quote_order.discounted:
    discounted_liner_cost = Order.get_single_liner_cost(quote_order)
    amount_discounted = quote_order.liner_cost - quote_order.discounted_single_liner_cost
    # Discount percentage
    extensions_cost.add_run("\nDiscount (" + str(discount_amount_percentage) + "%):                      "
                                                                               "                             "
                                                                               "                        "
                                                                               "             "
                                                                               "${:,.2f}".format(amount_discounted))
    # New cost of liner
    extensions_cost.add_run("\nNew cost of liner:                                            "
                            "                                           ${:,.2f}".format(discounted_liner_cost))

# Prints out total liner cost if added liners
if total_liners > 1:

    # Prints total discounted liner cost
    if discounted:
        discounted_total_liner_cost = total_liners * discounted_liner_cost
        extensions_cost.add_run("\nTotal cost of (" + str(total_liners) + ") liners:                             "
                                                                          "                                       "
                                                                          "        "
                                                                          "${:,.2f}".format(
            discounted_liner_cost * total_liners))
    # Prints total liner cost
    else:
        extensions_cost.add_run("\nTotal cost of (" + str(total_liners) + ") liners:                            "
                                                                          "                                     "
                                                                          "          "
                                                                          "${:,.2f}".format(
            Order.get_total_liners_cost(quote_order)))

# Check to see if customizations were ordered
if len(order_list) > 0:

    # Loop through order list to add to LINING SYSTEM
    for order in order_list:

        # Geo
        if order[0] == 'g' and order[1] == 'e':
            # Print out floor layer
            extensions_cost.add_run("\nTotal Geo cost:                                              "
                                    "                                               ${:,.2f}"
                                    .format(Order.get_single_layer_geo_cost(quote_order, "both")))

        # Batten strips
        elif order[0] == 'b' and order[1] == 'a':
            extensions_cost.add_run("\nBatten strips:                      "
                                    "                                                                           "
                                    "${:,.2f}".format(Order.get_batten_strip_cost(quote_order)))

        # J-bolts
        elif order[0] == 'j' and order[1] == 'b':
            extensions_cost.add_run("\n(" + str(Order.get_jbolt_number(quote_order)) + ") J-bolts:          "
                                                                                       "                    "
                                                                                       "                     "
                                                                                       "                         "
                                                                                       "                    "
                                                                                       "${:,.2f}".format(
                Order.get_jbolt_cost(quote_order)))

        # Lifting hem
        elif order[0] == 'l' and order[1] == 'i':
            extensions_cost.add_run("\nLifting hem:                                                                  "
                                    "                                   "
                                    "${:,.2f}".format(Order.get_lifting_hem_cost(quote_order)))

# Price of one lining system
lining_system_underline = extensions_cost.add_run(
    "\nTotal cost for lining system:                                                                   "
    "${:,.2f}".format(Order.get_single_lining_system(quote_order)))

# Calculate liner addition configuration
if total_liners > 1:
    total_lining_system_cost = Order.get_single_lining_system(quote_order) * total_liners
    lining_system_underline = extensions_cost.add_run("\nTotal cost for " + num2words(total_liners) + " ("
                                                      + str(
        total_liners) + ") lining systems:                              "
                        "                  ${:,.2f}".format(total_lining_system_cost))

# Pretty cosmetics for LINING SYSTEM subsection
lining_system_underline.underline = WD_UNDERLINE.SINGLE
extensions_cost.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# INSTALLATION section of quote

# See if necessary to include
if order_list.__contains__("installation"):

    # Prints out header for installation section
    installation_header_paragraph = quote.add_paragraph()
    installation_header_underline = installation_header_paragraph.add_run("\nInstallation for "
                                                                          + num2words(total_liners)
                                                                          + f" ({total_liners}) lining system")
    if total_liners > 1:
        installation_header_paragraph.add_run("s")
    installation_header_paragraph.add_run(":")

    # Paragraph formatting
    installation_header_underline.underline = True

    # Prints out single installation cost
    installation_paragraph = quote.add_paragraph("Installation:                                                    "
                                                 "                                            ${:,.2f}"
                                                 .format(installation_cost))

    # Single travel cost
    installation_paragraph.add_run("\nMobilization:                                                               "
                                   "                               ${:,.2f}".format(travel_cost))

    # Figures out total install package
    total_installation_package = (installation_cost * Order.get_total_liners(quote_order)) + travel_cost

    # Multiple liners
    if total_liners > 1:
        total_installation_package *= total_liners
        installation_paragraph.add_run(
            "\nSingle installation package:                                                            "
            "      ${:,.2f}".format(total_installation_cost))
    installation_paragraph.add_run(
        "\nTotal installation package:                                                           "
        "       ${:,.2f}".format(total_installation_package))

    # Paragraph formatting
    installation_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Calculations for next two sections

# Calculates total turnkey cost
turnkey_cost = 0
if total_liners > 1:
    turnkey_cost += total_lining_system_cost
else:
    turnkey_cost += Order.get_single_lining_system(quote_order)
if order_list.__contains__("installation"):
    turnkey_cost += total_installation_package

# Calculates taxes to charge
tax_cost = (tax_perc / 100) * turnkey_cost
total_turnkey_cost = turnkey_cost + tax_cost

# TAXES section of quote

if tax_perc > 0:
    taxes_paragraph = quote.add_paragraph()

    # Amount of turnkey cost taxed
    taxes_paragraph.add_run("\nAmount taxed (" + str(tax_perc) + "%):                            "
                                                                "                                "
                                                                "              ${:,.2f}".format(tax_cost))

# INSTALLATION wrap up section of quote

installation_wrap_up_paragraph = quote.add_paragraph()

# Prints out total turnkey cost
turnkey_underline = installation_wrap_up_paragraph.add_run("\n\nTotal Turnkey Cost:                           "
                                                           "                                                   "
                                                           "   ${:,.2f}".format(total_turnkey_cost))
turnkey_underline.underline = WD_UNDERLINE.SINGLE

# Paragraph formatting
installation_wrap_up_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Save quote document as quote
quote.save("Quote.docx")
